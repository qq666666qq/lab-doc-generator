"""
组装 Word 文档
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_docx(
    steps: list[dict],
    screenshots: list[str],
    user: str = "root",
    host: str = "server",
    output_path: str = "output/lab_report.docx",
) -> str:
    """
    组装 Word 文档

    Args:
        steps: 步骤列表 [{step, title, command}]
        screenshots: 截图文件路径列表（与 steps 对应）
        user: 用户名
        host: 主机名
        output_path: 输出文件路径

    Returns:
        生成的 Word 文件路径
    """
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

    doc = Document()

    # 标题
    title = doc.add_heading("实验操作记录", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 终端信息
    info_para = doc.add_paragraph()
    info_run = info_para.add_run(f"终端环境：{user}@{host}")
    info_run.font.size = Pt(11)
    info_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # 空行

    # 遍历步骤
    for i, step in enumerate(steps):
        # 步骤标题
        heading = doc.add_heading(
            f"步骤 {step['step']}：{step['title']}", level=2
        )

        # 命令行（终端风格）
        prompt = f"{user}@{host}:~$"
        cmd_para = doc.add_paragraph()
        cmd_para.paragraph_format.space_before = Pt(4)
        cmd_para.paragraph_format.space_after = Pt(2)
        # prompt 部分
        prompt_run = cmd_para.add_run(prompt + " ")
        prompt_run.font.name = "Courier New"
        prompt_run.font.size = Pt(10)
        prompt_run.font.color.rgb = RGBColor(39, 150, 63)
        # 命令部分
        cmd_run = cmd_para.add_run(step["command"])
        cmd_run.font.name = "Courier New"
        cmd_run.font.size = Pt(10)

        # 截图
        if i < len(screenshots) and os.path.exists(screenshots[i]):
            try:
                doc.add_picture(screenshots[i], width=Inches(5.5))
                # 居中图片
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[截图加载失败: {e}]")

        # 步骤间空行
        doc.add_paragraph()

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    # 测试用
    result = build_docx(
        steps=[
            {"step": 1, "title": "测试", "command": "echo hello"},
        ],
        screenshots=[],
    )
    print(f"文档已保存: {result}")
